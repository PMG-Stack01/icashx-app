from docx import Document
import os

def generate_purchase_contract(name, address, price, closing_date):
    doc = Document()
    doc.add_heading("Purchase Agreement", 0)
    doc.add_paragraph(f"Seller: {name}")
    doc.add_paragraph(f"Property Address: {address}")
    doc.add_paragraph(f"Purchase Price: ${price}")
    doc.add_paragraph(f"Closing Date: {closing_date}")
    file_path = f"/home/user/contracts/{name}_contract.docx"
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)
    return file_path